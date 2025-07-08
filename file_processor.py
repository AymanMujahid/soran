import os
import asyncio
import aiofiles
import mimetypes
from pathlib import Path
from typing import List, Dict, Optional, Tuple, Any, AsyncGenerator
from datetime import datetime
import hashlib
import logging

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
import PyPDF2
from PIL import Image
from sqlalchemy.orm import Session

from vector_database import VectorChunk, create_text_hash
from embeddings import get_embedding_service

logger = logging.getLogger(__name__)

class ProcessingError(Exception):
    pass

class FileTypeError(Exception):
    pass

class FileSizeError(Exception):
    pass

class DocumentChunk:
    def __init__(self, text: str, metadata: Dict[str, Any] = None, chunk_index: int = 0):
        self.text = text.strip()
        self.metadata = metadata or {}
        self.chunk_index = chunk_index
        self.id = create_text_hash(f"{text}_{chunk_index}")
        self.created_at = datetime.utcnow()

class FileProcessor:
    """Base class for file processors"""
    
    SUPPORTED_EXTENSIONS = []
    
    def __init__(self):
        self.chunk_size = 1000
        self.chunk_overlap = 100
    
    async def can_process(self, file_path: Path) -> bool:
        """Check if file can be processed"""
        return file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    async def process(self, file_path: Path, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """Process file and return chunks"""
        raise NotImplementedError
    
    def _split_text_into_chunks(self, text: str, max_chunk_size: int = None, overlap: int = None) -> List[str]:
        """Split text into overlapping chunks"""
        if not text:
            return []
        
        max_size = max_chunk_size or self.chunk_size
        overlap_size = overlap or self.chunk_overlap
        
        if len(text) <= max_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_size
            
            if end >= len(text):
                chunks.append(text[start:])
                break
            
            # Try to end at a sentence boundary
            chunk_end = text.rfind('.', start, end)
            if chunk_end == -1:
                # Try to end at a word boundary
                chunk_end = text.rfind(' ', start, end)
            
            if chunk_end == -1 or chunk_end <= start:
                chunk_end = end
            
            chunks.append(text[start:chunk_end])
            start = chunk_end - overlap_size
            
            if start <= 0:
                start = chunk_end
        
        return [chunk.strip() for chunk in chunks if chunk.strip()]

class PDFProcessor(FileProcessor):
    SUPPORTED_EXTENSIONS = ['.pdf']
    
    async def process(self, file_path: Path, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """Process PDF file"""
        try:
            logger.info(f"Processing PDF: {file_path}")
            
            text_chunks = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                total_pages = len(pdf_reader.pages)
                full_text = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            full_text.append(page_text)
                            logger.debug(f"Extracted text from page {page_num + 1}/{total_pages}")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
                
                if full_text:
                    combined_text = '\n'.join(full_text)
                    chunks = self._split_text_into_chunks(combined_text)
                    
                    for i, chunk_text in enumerate(chunks):
                        chunk_metadata = {
                            'file_type': 'pdf',
                            'total_pages': total_pages,
                            'chunk_index': i,
                            'total_chunks': len(chunks),
                            **(metadata or {})
                        }
                        text_chunks.append(DocumentChunk(chunk_text, chunk_metadata, i))
            
            logger.info(f"PDF processing complete: {len(text_chunks)} chunks generated")
            return text_chunks
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise ProcessingError(f"PDF processing failed: {e}")

class DOCXProcessor(FileProcessor):
    SUPPORTED_EXTENSIONS = ['.docx']
    
    async def process(self, file_path: Path, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """Process DOCX file"""
        try:
            logger.info(f"Processing DOCX: {file_path}")
            
            doc = Document(file_path)
            
            paragraphs = []
            tables_text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        table_data.append(' | '.join(row_data))
                
                if table_data:
                    tables_text.append('\n'.join(table_data))
            
            # Combine all text
            all_text = []
            if paragraphs:
                all_text.extend(paragraphs)
            if tables_text:
                all_text.extend(tables_text)
            
            combined_text = '\n'.join(all_text)
            chunks = self._split_text_into_chunks(combined_text)
            
            text_chunks = []
            for i, chunk_text in enumerate(chunks):
                chunk_metadata = {
                    'file_type': 'docx',
                    'total_paragraphs': len(paragraphs),
                    'total_tables': len(doc.tables),
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                    **(metadata or {})
                }
                text_chunks.append(DocumentChunk(chunk_text, chunk_metadata, i))
            
            logger.info(f"DOCX processing complete: {len(text_chunks)} chunks generated")
            return text_chunks
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise ProcessingError(f"DOCX processing failed: {e}")

class ExcelProcessor(FileProcessor):
    SUPPORTED_EXTENSIONS = ['.xlsx', '.xls']
    
    async def process(self, file_path: Path, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """Process Excel file"""
        try:
            logger.info(f"Processing Excel: {file_path}")
            
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None)
            
            text_chunks = []
            
            for sheet_name, df in excel_data.items():
                # Convert DataFrame to text
                sheet_text = f"Sheet: {sheet_name}\n"
                
                # Add column headers
                sheet_text += "Columns: " + ", ".join(df.columns.astype(str)) + "\n\n"
                
                # Add data rows
                for index, row in df.iterrows():
                    row_text = " | ".join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])
                    if row_text:
                        sheet_text += row_text + "\n"
                
                # Split sheet into chunks
                chunks = self._split_text_into_chunks(sheet_text, max_chunk_size=1500)
                
                for i, chunk_text in enumerate(chunks):
                    chunk_metadata = {
                        'file_type': 'excel',
                        'sheet_name': sheet_name,
                        'total_sheets': len(excel_data),
                        'sheet_rows': len(df),
                        'sheet_columns': len(df.columns),
                        'chunk_index': i,
                        'total_chunks': len(chunks),
                        **(metadata or {})
                    }
                    text_chunks.append(DocumentChunk(chunk_text, chunk_metadata, i))
            
            logger.info(f"Excel processing complete: {len(text_chunks)} chunks generated")
            return text_chunks
            
        except Exception as e:
            logger.error(f"Error processing Excel {file_path}: {e}")
            raise ProcessingError(f"Excel processing failed: {e}")

class TextProcessor(FileProcessor):
    SUPPORTED_EXTENSIONS = ['.txt', '.md', '.csv']
    
    async def process(self, file_path: Path, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """Process text file"""
        try:
            logger.info(f"Processing text file: {file_path}")
            
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = await f.read()
            
            if not content.strip():
                return []
            
            # Special handling for CSV
            if file_path.suffix.lower() == '.csv':
                try:
                    df = pd.read_csv(file_path)
                    content = df.to_string(index=False)
                except:
                    pass  # Fall back to regular text processing
            
            chunks = self._split_text_into_chunks(content)
            
            text_chunks = []
            for i, chunk_text in enumerate(chunks):
                chunk_metadata = {
                    'file_type': 'text',
                    'file_extension': file_path.suffix.lower(),
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                    'character_count': len(chunk_text),
                    **(metadata or {})
                }
                text_chunks.append(DocumentChunk(chunk_text, chunk_metadata, i))
            
            logger.info(f"Text processing complete: {len(text_chunks)} chunks generated")
            return text_chunks
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            raise ProcessingError(f"Text processing failed: {e}")

class HTMLProcessor(FileProcessor):
    SUPPORTED_EXTENSIONS = ['.html', '.htm']
    
    async def process(self, file_path: Path, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """Process HTML file"""
        try:
            logger.info(f"Processing HTML: {file_path}")
            
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = await f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer"]):
                script.decompose()
            
            # Extract text from various elements
            text_elements = []
            
            for element in soup.find_all(['p', 'div', 'article', 'section', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li']):
                text = element.get_text(strip=True)
                if text and len(text) > 20:
                    text_elements.append(text)
            
            combined_text = '\n'.join(text_elements)
            chunks = self._split_text_into_chunks(combined_text)
            
            text_chunks = []
            for i, chunk_text in enumerate(chunks):
                chunk_metadata = {
                    'file_type': 'html',
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                    'total_elements': len(text_elements),
                    **(metadata or {})
                }
                text_chunks.append(DocumentChunk(chunk_text, chunk_metadata, i))
            
            logger.info(f"HTML processing complete: {len(text_chunks)} chunks generated")
            return text_chunks
            
        except Exception as e:
            logger.error(f"Error processing HTML {file_path}: {e}")
            raise ProcessingError(f"HTML processing failed: {e}")

class FileProcessingService:
    """Main service for processing files"""
    
    def __init__(self, max_file_size: int = 50 * 1024 * 1024):  # 50MB
        self.max_file_size = max_file_size
        self.processors = {
            'pdf': PDFProcessor(),
            'docx': DOCXProcessor(),
            'excel': ExcelProcessor(),
            'text': TextProcessor(),
            'html': HTMLProcessor()
        }
        
        # Build extension to processor mapping
        self.extension_map = {}
        for processor_name, processor in self.processors.items():
            for ext in processor.SUPPORTED_EXTENSIONS:
                self.extension_map[ext] = processor_name
        
        logger.info(f"File processing service initialized with {len(self.processors)} processors")
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return list(self.extension_map.keys())
    
    def _validate_file(self, file_path: Path) -> Tuple[str, str]:
        """Validate file and return file type and mime type"""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            raise FileSizeError(f"File size {file_size} exceeds maximum {self.max_file_size}")
        
        file_ext = file_path.suffix.lower()
        if file_ext not in self.extension_map:
            raise FileTypeError(f"Unsupported file type: {file_ext}")
        
        mime_type, _ = mimetypes.guess_type(str(file_path))
        if not mime_type:
            mime_type = "application/octet-stream"
        
        return self.extension_map[file_ext], mime_type
    
    async def process_file(self, file_path: Path, additional_metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """Process a file and return document chunks"""
        try:
            processor_type, mime_type = self._validate_file(file_path)
            
            base_metadata = {
                'file_name': file_path.name,
                'file_path': str(file_path),
                'file_size': file_path.stat().st_size,
                'mime_type': mime_type,
                'processed_at': datetime.utcnow().isoformat(),
                'processor_type': processor_type
            }
            
            if additional_metadata:
                base_metadata.update(additional_metadata)
            
            processor = self.processors[processor_type]
            chunks = await processor.process(file_path, base_metadata)
            
            logger.info(f"Successfully processed {file_path.name}: {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}")
            raise
    
    async def process_files_batch(self, file_paths: List[Path], max_concurrent: int = 3) -> Dict[str, Any]:
        """Process multiple files concurrently"""
        results = {
            'successful': {},
            'failed': {},
            'total_chunks': 0
        }
        
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def process_single_file(file_path: Path):
            async with semaphore:
                try:
                    chunks = await self.process_file(file_path)
                    results['successful'][str(file_path)] = {
                        'chunks': len(chunks),
                        'chunks_data': chunks
                    }
                    results['total_chunks'] += len(chunks)
                    logger.info(f"✅ Processed {file_path.name}: {len(chunks)} chunks")
                except Exception as e:
                    results['failed'][str(file_path)] = str(e)
                    logger.error(f"❌ Failed to process {file_path.name}: {e}")
        
        tasks = [process_single_file(file_path) for file_path in file_paths]
        await asyncio.gather(*tasks, return_exceptions=True)
        
        logger.info(f"Batch processing complete: {len(results['successful'])} successful, {len(results['failed'])} failed")
        return results
    
    async def process_and_store(self, file_path: Path, bot_id: str, embedding_model: str = "all-MiniLM-L6-v2") -> Dict[str, Any]:
        """Process file and store in vector database"""
        try:
            from vector_database import get_vector_database
            
            # Process file
            chunks = await self.process_file(file_path)
            
            if not chunks:
                return {'success': False, 'error': 'No content extracted from file'}
            
            # Get embedding service and vector database
            embedding_service = get_embedding_service(embedding_model)
            vector_db = get_vector_database(bot_id, embedding_service.get_dimension())
            
            # Generate embeddings for chunks
            texts = [chunk.text for chunk in chunks]
            embeddings = embedding_service.encode_texts(texts)
            
            # Create vector chunks
            vector_chunks = []
            for chunk, embedding in zip(chunks, embeddings):
                vector_chunk = VectorChunk(
                    id=chunk.id,
                    text=chunk.text,
                    embedding=embedding,
                    metadata=chunk.metadata
                )
                vector_chunks.append(vector_chunk)
            
            # Store in vector database
            success = vector_db.add_vectors(vector_chunks)
            
            if success:
                result = {
                    'success': True,
                    'file_name': file_path.name,
                    'chunks_processed': len(chunks),
                    'chunks_stored': len(vector_chunks),
                    'bot_id': bot_id,
                    'embedding_model': embedding_model
                }
                logger.info(f"✅ File processed and stored: {file_path.name} - {len(chunks)} chunks")
                return result
            else:
                return {'success': False, 'error': 'Failed to store vectors'}
                
        except Exception as e:
            logger.error(f"Error in process_and_store for {file_path}: {e}")
            return {'success': False, 'error': str(e)}

# Global file processing service
file_processing_service = FileProcessingService()

def get_file_processing_service() -> FileProcessingService:
    """Get the global file processing service"""
    return file_processing_service